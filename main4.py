####################################################################################################
#                                      Main Long Term                                              #
####################################################################################################

import six
import math
import time
import Quandl
import calendar
import warnings
import numpy as np
import pandas as pd
import seaborn as sb
import pylab as pylab
from docx import Document
from datetime import datetime

from matplotlib import colors
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from mpl_toolkits.mplot3d import Axes3D

from scipy.interpolate import interp1d
import statsmodels.stats.stattools as stats_stattools
import statsmodels.tsa.stattools as tsa_stattools
import statsmodels.tsa.seasonal as tsa_seasonal
import statsmodels.api as sm
import xgboost as xgb
from unbalanced_dataset import SMOTE

from sklearn import svm
from sklearn import metrics, cross_validation, linear_model, naive_bayes, neighbors, ensemble
from sklearn import feature_selection
from sklearn import decomposition
from sklearn import discriminant_analysis
from sklearn import preprocessing

import sys
from os import listdir
from os.path import isfile, join
from helpers import features_analysis, procces_stocks, data_manipulation, download_quandl_data, ml_dataset, classifier_utils, report_generator, Iteration, Stacking, Boosting

fig_size = [10, 6]
plt.rcParams["figure.figsize"] = fig_size
sb.set_style('darkgrid')
warnings.filterwarnings("ignore", category=DeprecationWarning)


GOLD = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/GOLD.csv')
SILVER = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/SILVER.csv')
PLAT = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/PLAT.csv')
OIL_BRENT = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/OIL_BRENT.csv')

USD_GBP = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/USD_GBP.csv')
JPY_USD = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/JPY_USD.csv')
AUD_USD = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/AUD_USD.csv')

INDEX_DJIA = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_DJIA.csv')
INDEX_HSI = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_HSI.csv')
INDEX_IBEX = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_IBEX.csv')
INDEX_N225 = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_N225.csv')
INDEX_SP500 = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_SP500.csv')
INDEX_AXJO = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_AXJO.csv')
INDEX_FCHI = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_FCHI.csv')
INDEX_GDAXI = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_GDAXI.csv')

values_names = ['GOLD', 'SILVER', 'PLAT', 'OIL_BRENT', 'USD_GBP', 'JPY_USD', 'AUD_USD', 'DJIA', 'HSI', 'IBEX', 'N225', 'SP500', 'AXJO', 'FCHI', 'GDAXI']
values_dfs = [GOLD, SILVER, PLAT, OIL_BRENT, USD_GBP, JPY_USD, AUD_USD, INDEX_DJIA, INDEX_HSI, INDEX_IBEX, INDEX_N225, INDEX_SP500, INDEX_AXJO, INDEX_FCHI, INDEX_GDAXI]
values_cols = ['USD', 'Value', 'Open', 'Close', 'High', 'Low', 'Volume']
dict_dfs_cols = {}

for index in range(len(values_names)):
    name = values_names[index]
    df = values_dfs[index]
    cols = df.columns.values
    #new_cols = [x for x in cols if x not in ['Date']]
    new_cols = [x for x in cols if x not in ['Date', 'USD', 'Value', 'Open', 'Close', 'High', 'Low', 'Volume']]
    
    dict_dfs_cols[name] = new_cols

dataset = ml_dataset.generate_df_dataset(values_names, values_dfs, dict_dfs_cols)

#First 30 row
dataset = dataset[31:]
dataset = dataset.reset_index(drop=True)
datasetY = dataset.copy(deep=True)
#dataset = dataset.fillna(method='ffill')


SEED = 42
algorithms_list = ["", "LRC", "SGD", "LDA", "QDA", "SVM", "NBG", "NBB", "KNN", "ABC", "BGC", "GBC", "RFC", "ETC"]

document = Document()
document.add_heading('Experimentos', 0)

## Table headings
table = document.add_table(rows=1, cols=len(algorithms_list))
hdr_cells = table.rows[0].cells
for i in range(len(algorithms_list)):
    hdr_cells[i].text = algorithms_list[i]
    

for colY, regex in zip(['IBEX_RD_B1_Close', 'IBEX_RD_B2_Close', 'IBEX_RD_B3_Close', 'IBEX_RD_B5_Close', 'IBEX_RD_B10_Close', 'IBEX_RD_B20_Close', 'IBEX_RD_B30_Close', 'IBEX_RD_B50_Close', 'IBEX_RD_B100_Close'],
                       ['RD_B1_Close', 'RD_B2_Close', 'RD_B3_Close', 'RD_B5_Close', 'RD_B10_Close', 'RD_B20_Close', 'RD_B30_Close', 'RD_B50_Close', 'RD_B100_Close']): 
    print "##############################"
    print "## Predicintg %s , regex %s ##" % (colY, regex)
    models_score = []
    
    for model in [linear_model.LogisticRegression(),linear_model.SGDClassifier(), 
                  discriminant_analysis.LinearDiscriminantAnalysis(), discriminant_analysis.QuadraticDiscriminantAnalysis(),
                  svm.SVC(),
                  naive_bayes.GaussianNB(), naive_bayes.BernoulliNB(),
                  neighbors.KNeighborsClassifier(),
                  ensemble.AdaBoostClassifier(), ensemble.BaggingClassifier(), ensemble.GradientBoostingClassifier(), ensemble.RandomForestClassifier(), ensemble.ExtraTreesClassifier()]:


        colsToShift = 1

        df_x = dataset.filter(regex=(regex))
        last_row = list(range(df_x.shape[0] - colsToShift, df_x.shape[0] ))
        df_x = df_x.drop(last_row, axis=0)
        df_x = df_x.drop(colY, axis=1)

        df_y = dataset[colY].shift(-colsToShift)
        last_row = list(range(df_y.shape[0] - colsToShift, df_y.shape[0] ))
        df_y = df_y.drop(last_row, axis=0)

        training_dates = Iteration.Iteration('2008-06-17', '2011-09-01')
        testing_dates  = Iteration.Iteration('2012-09-04', '2014-10-06')
        training_dates.calculate_indices(dataset)
        testing_dates.calculate_indices(dataset)

        trainDates = []
        testDates = []
        trainDates.append(training_dates.lowerIndex)
        trainDates.append(training_dates.upperIndex)
        testDates.append(testing_dates.lowerIndex)
        testDates.append(testing_dates.upperIndex)

        total = (trainDates[1]-trainDates[0]) + (testDates[1]-testDates[0])
        tr = float(trainDates[1]-trainDates[0]) / total * 100.0
        te = float(testDates[1]-testDates[0]) / total * 100.0

        trainX, trainY, testX, testY = ml_dataset.train_arrays_experiments(df_x, df_y, trainDates, testDates)
        model.fit(trainX, trainY)
        y_pred = model.predict(testX)  
        score = metrics.accuracy_score(testY, y_pred)
        models_score.append(score)
        #print "Score %s \t %s" % (score, model.__class__.__name__)
        
    ##  header
    row_cells = table.add_row().cells
        
    row_cells[0].text = str(regex)
    col = 1

    ##Table test X_X row
    for model in range(len(models_score)):
        cell = ("%.2f%%\n" % (models_score[model]*100))
        row_cells[col].text = cell
        col += 1
document.save('experiments.docx')


## Predicions  ##
for i in range(testY.shape[0]):
    print "## TODAY - %s ##" % dataset_all.iloc[testing_dates.lowerIndex+i]['Date']
    print "Date %s \t Close %f \t Bin %d" % (dataset_all.iloc[testing_dates.lowerIndex-9+i]['Date'], dataset_all.iloc[testing_dates.lowerIndex-9+i]['IBEX_Close'],dataset_all.iloc[testing_dates.lowerIndex-9+i]['IBEX_RD_B10_Close']) 
    print "Date %s \t Close %f \t Bin %d" % (dataset_all.iloc[testing_dates.lowerIndex+i]['Date'], dataset_all.iloc[testing_dates.lowerIndex+i]['IBEX_Close'],dataset_all.iloc[testing_dates.lowerIndex+i]['IBEX_RD_B10_Close'])
    print "Date %s \t Close %f \t Bin %d" % (dataset_all.iloc[testing_dates.lowerIndex+1+i]['Date'], dataset_all.iloc[testing_dates.lowerIndex+1+i]['IBEX_Close'], dataset_all.iloc[testing_dates.lowerIndex+1+i]['IBEX_RD_B10_Close'])
    
    print "Test %d \t Pred %d" % (testY[i], y_pred[i])
    if testY[i] == 1:
        print "El índice SUBE de precio desde el día %s hasta el día %s" % (dataset_all.iloc[testing_dates.lowerIndex-9+i]['Date'],dataset_all.iloc[testing_dates.lowerIndex+1+i]['Date'])
        print "* BUY shares for day %s *" % dataset_all.iloc[testing_dates.lowerIndex+1+i]['Date']
        #Buy only if today`s price is lower than 10 days ago
    else:
        print "El índice BAJA de precio desde el día %s hasta el día %s" % (dataset_all.iloc[testing_dates.lowerIndex-9+i]['Date'],dataset_all.iloc[testing_dates.lowerIndex+1+i]['Date'])        
        print "* SELL shares for day %s *" % dataset_all.iloc[testing_dates.lowerIndex+1+i]['Date']