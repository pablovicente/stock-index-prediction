####################################################################################################
#                                      Main Type 6                                                 #
####################################################################################################

import six
import math
import time
import Quandl
import calendar
import warnings
import numpy as np
import pandas as pd
import seaborn as sb
import pylab as pylab
from docx import Document
from datetime import datetime

from matplotlib import colors
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from mpl_toolkits.mplot3d import Axes3D

from scipy.interpolate import interp1d
import statsmodels.stats.stattools as stats_stattools
import statsmodels.tsa.stattools as tsa_stattools
import statsmodels.tsa.seasonal as tsa_seasonal
import statsmodels.api as sm
import xgboost as xgb
from unbalanced_dataset import SMOTE

from sklearn import svm
from sklearn import metrics, cross_validation, linear_model, naive_bayes, neighbors, ensemble
from sklearn import feature_selection
from sklearn import decomposition
from sklearn import discriminant_analysis
from sklearn import preprocessing

import sys
from os import listdir
from os.path import isfile, join
from helpers import features_analysis, procces_stocks, data_manipulation, download_quandl_data, ml_dataset, classifier_utils, report_generator, Iteration, Stacking, Boosting

fig_size = [10, 6]
plt.rcParams["figure.figsize"] = fig_size
sb.set_style('darkgrid')
warnings.filterwarnings("ignore", category=DeprecationWarning)


GOLD = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/GOLD.csv')
SILVER = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/SILVER.csv')
PLAT = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/PLAT.csv')
OIL_BRENT = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/OIL_BRENT.csv')

USD_GBP = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/USD_GBP.csv')
JPY_USD = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/JPY_USD.csv')
AUD_USD = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/AUD_USD.csv')

INDEX_DJIA = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_DJIA.csv')
INDEX_HSI = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_HSI.csv')
INDEX_IBEX = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_IBEX.csv')
INDEX_N225 = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_N225.csv')
INDEX_SP500 = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_SP500.csv')
INDEX_AXJO = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_AXJO.csv')
INDEX_FCHI = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_FCHI.csv')
INDEX_GDAXI = data_manipulation.read_csv_data('/Users/Pablo/Desktop/TFM/Data/INDEX_GDAXI.csv')

values_names = ['GOLD', 'SILVER', 'PLAT', 'OIL_BRENT', 'USD_GBP', 'JPY_USD', 'AUD_USD', 'DJIA', 'HSI', 'IBEX', 'N225', 'SP500', 'AXJO', 'FCHI', 'GDAXI']
values_dfs = [GOLD, SILVER, PLAT, OIL_BRENT, USD_GBP, JPY_USD, AUD_USD, INDEX_DJIA, INDEX_HSI, INDEX_IBEX, INDEX_N225, INDEX_SP500, INDEX_AXJO, INDEX_FCHI, INDEX_GDAXI]
values_cols = ['USD', 'Value', 'Open', 'Close', 'High', 'Low', 'Volume']
dict_dfs_cols = {}

for index in range(len(values_names)):
    name = values_names[index]
    df = values_dfs[index]
    cols = df.columns.values
    #new_cols = [x for x in cols if x not in ['Date']]
    new_cols = [x for x in cols if x not in ['Date', 'USD', 'Value', 'Open', 'Close', 'High', 'Low', 'Volume']]
    
    dict_dfs_cols[name] = new_cols

dataset = ml_dataset.generate_df_dataset(values_names, values_dfs, dict_dfs_cols)

#First 30 row
dataset = dataset[31:]
dataset = dataset.reset_index(drop=True)
datasetY = dataset.copy(deep=True)
#dataset = dataset.fillna(method='ffill')



training_dates = Iteration.Iteration('2009-08-19', '2014-12-01')
testing_dates  = Iteration.Iteration('2014-12-02', '2016-04-20')
training_dates.calculate_indices(dataset)
testing_dates.calculate_indices(dataset)

trainDates = []
testDates = []
trainDates.append(training_dates.lowerIndex)
trainDates.append(training_dates.upperIndex)
testDates.append(testing_dates.lowerIndex)
testDates.append(testing_dates.upperIndex)

trainX, trainY, testX, testY, cols = ml_dataset.dataset_to_train_using_dates(dataset, trainDates, testDates, binary=False, shiftFeatures=False, shiftTarget=False)

N_TREES = 500
SEED = 42
log = open("log.txt", "w")

selected_models = [
    "LRC:all_greedy",   
    "SGD:all_greedy",

#    "LDA:all_greedy",
#    "QDA:all_greedy",
#   
#    "SVM:all_greedy",
#   
#    "NBG:all_greedy",
#    "NBB:all_greedy",
#
#    "KNN:all_greedy", 
#    
#    "ABC:all_greedy",
#    "BGC:all_greedy",
#    "GBC:all_greedy", 
#    "RFC:all_greedy",
    "ETC:all_greedy"
]

algorithms_list = ['']

# Create the models on the fly
models = []
for item in selected_models:
    model_id, data_set = item.split(':')
    model = {'LRC':linear_model.LogisticRegression,
             'SGD':linear_model.SGDClassifier,
             
             'LDA':discriminant_analysis.LinearDiscriminantAnalysis,
             'QDA':discriminant_analysis.QuadraticDiscriminantAnalysis,
             
             'SVM':svm.SVC,
            
             'NBG':naive_bayes.GaussianNB,
             'NBB':naive_bayes.BernoulliNB,
             
             "KNN":neighbors.KNeighborsClassifier,

             'ABC': ensemble.AdaBoostClassifier,
             'BGC': ensemble.BaggingClassifier,             
             'GBC': ensemble.GradientBoostingClassifier,
             'RFC': ensemble.RandomForestClassifier,
             'ETC': ensemble.ExtraTreesClassifier        
            }[model_id]()
    models.append((model, data_set))
    algorithms_list.append(model_id)


grid_search = True
## Set params
for model, feature_set in models:
    model.set_params(**classifier_utils.find_params(model, feature_set, trainX, trainY, grid_search))

colY = 'IBEX_RD_B1_Close'
experiments = {'Test 1': ["Log_Return_1_Close", "Log_Return_.*_Close"],#, "RD1_Close", "RD.?_Close", "RD_P1_Close", "RD_P.?_Close", 
#                          "RD_B1_Close", "RD_B.?_Close"],
#               'Test 2': ["RD1_(Open|Close|High|Low|Volume)", "RD.?_(Open|Close|High|Low|Volume)", "RD_P1_(Open|Close|High|Low|Volume)", 
#                          "RD_P.?_(Open|Close|High|Low|Volume)", "RD_B1_(Open|Close|High|Low|Volume)", "RD_B.?_(Open|Close|High|Low|Volume)"],
#               'Test 3': ["Log_Return_1_USD", "Log_Return_.*_USD", "^(GOLD|SILVER|PLAT|OIL_BRENT)_RD1", "^(GOLD|SILVER|PLAT|OIL_BRENT)_RD.?", 
#                          "^(GOLD|SILVER|PLAT|OIL_BRENT)_RD_P1", "^(GOLD|SILVER|PLAT|OIL_BRENT)_RD_P.?", "^(GOLD|SILVER|PLAT|OIL_BRENT)_RD_B1", 
#                          "^(GOLD|SILVER|PLAT|OIL_BRENT)_RD_B.?"],
#               'Test 4': ["Log_Return_1_(USD|Close)", "Log_Return_.*_(USD|Close)","RD1_(USD|Close)", "RD.?_(USD|Close)", "RD_P1_(USD|Close)", 
#                          "RD_P.?_(USD|Close)", "RD_B1_(USD|Close)", "RD_B.?_(USD|Close)"],
               'Test 5': ["Log_Return_1", "Log_Return_.*"]#, "RD1", "RD_P.?_", "RD_P1", "RD_P.?_", "RD_B1", "RD_B.?_", ".*"]
              }    

trainingPeriodList = [(365,'one_year'), (365*2,'two_years'),(365*3,'three_years')]
testingPeriodList = [(1, '1_days'),(3, '3_days'),(5, '5_days'),(10, '10_days'), (40, '40_days')]
step = 365*3

dict_results = {'one_year': {},
                'two_years': {},
                'three_years': {}
               }

keys_exp = experiments.keys()
keys_exp.sort()
keys_temp = dict_results.keys()
keys_temp.sort()

for key in keys_temp:
    for test in keys_exp:
        dict_results[key][test] = {}     
        for num, days in testingPeriodList:
            dict_results[key][test][days] = {}
            for index in range(len(experiments[test])):
                dict_results[key][test][days][str(index+1)] = []

             
features_list = []
num_experiments = 0

#df_y = dataset[colY]
df_y = dataset[colY].shift(-1)
last_row = df_y.shape[0]-1
df_y = df_y.drop(last_row, axis=0)
print >> log, 'Dataset shape %s' % str(dataset.shape)
print >> log, 'Y shape %s' % str(df_y.shape)

num_tests = 0
start = time.time()

for trainingPeriod, key_year in trainingPeriodList:    
    
    print "=================================================="
    print "                    %s                     " % key_year 
    print "=================================================="
    print >> log, "=================================================="
    print >> log, "                    %s                     " % str(key_year)
    print >> log, "=================================================="

    for key_exp in keys_exp:
        regex_list = experiments[key_exp]
        sub_index = 1    

        for regex_element in regex_list:

            print "++++++++++++++++++++++++++++++++++++++++++++++++++"
            print "                    %s                  " % (key_exp + '-'+ str(sub_index)) 
            print "++++++++++++++++++++++++++++++++++++++++++++++++++"
            print >> log, "++++++++++++++++++++++++++++++++++++++++++++++++++"
            print >> log, "                   %s                  " % (key_exp + '.'+ str(sub_index))
            print >> log, "++++++++++++++++++++++++++++++++++++++++++++++++++"
        
            ##Trainig testing arrays
            df_x = dataset.filter(regex=(regex_element))
            df_x = df_x.drop(last_row, axis=0)
            iteration_ = 1

            for index in range(0,dataset.shape[0] - trainingPeriod - 41, step):
                print "--------------------------------------------------"
                print "                   Iteration %s           " % str(iteration_) 
                print "--------------------------------------------------"                
                print >> log, "--------------------------------------------------"                
                print >> log, "                   Iteration %s            " % str(iteration_)
                print >> log, "--------------------------------------------------"  

                trainDates = []
                trainDates.append(index)
                trainDates.append(index + trainingPeriod - 1)
                trainX, trainY = ml_dataset.only_train_array(df_x, df_y, trainDates)

                ## Stacking init
                clf = Stacking.Stacking(models, stack=False, fwls=False, model_selection=False, log=log)                
                ## Fit stacking model
                clf.fit(trainY, trainX)
                num_tests += 1
                iteration_ += 1
                
                for testingPeriod, key_test in testingPeriodList:
                
                        testDates = []
                        testDates.append(index + trainingPeriod)
                        testDates.append(index + trainingPeriod + testingPeriod)

                        total = (trainDates[1]-trainDates[0]) + (testDates[1]-testDates[0])
                        tr = float(trainDates[1]-trainDates[0]) / total * 100.0
                        te = float(testDates[1]-testDates[0]) / total * 100.0

                        print >> log, "Training: from %s to %s" % (str(dataset.Date[trainDates[0]]), str(dataset.Date[trainDates[1]]))
                        print >> log, "Testing: from %s to %s" % (str(dataset.Date[testDates[0]]), str(dataset.Date[testDates[1]]))
                        print >> log, "%.3f %% training %.3f %% testing" % (tr,te)
                        print >> log, "%d training %d testing" % (trainDates[1]-trainDates[0], testDates[1]-testDates[0])        

                        testX, testY = ml_dataset.only_test_array(df_x, df_y, testDates)
  
                        ###  Metrics
                        print >> log, "computing cv score"
                        mean_auc = 0.0
                        mean_accuracy = 0.0
                        iter_ = 1

                        cv_preds, models_score, models_f1 = clf.predict(trainY, trainX, testX, testY, show_steps=True)
                        cv_preds_bin = np.round_(cv_preds, decimals=0)
                        accuracy = metrics.accuracy_score(testY, cv_preds_bin)
                        f1 = metrics.f1_score(testY, cv_preds_bin)
                        #print "Accuracy: %.2f" % accuracy


                        
                        #for i in range(len(models_score)):                            
                        #    dict_results[key_year][key_exp][str(sub_index)][i] += models_score[i]*100
                        dict_results[key_year][key_exp][key_test][str(sub_index)].append(models_score)
                        #dict_results2[key_year][key_exp][str(sub_index)].append(models_score)

            
            sub_index += 1       


document = Document()
document.add_heading('Experimentos', 0)

keys_exp = experiments.keys()
keys_exp.sort()


keys_temp = dict_results.keys()
for key in keys_temp:
    print "=================="
    #print "new dates %s" % key
    document.add_heading("Dates  " + str(key), level=1)
    for testingPeriod, days in testingPeriodList: 
        #print "---------------------"        
        #print "new days %s" % days
        document.add_heading("Days  " + str(days), level=1)

        ## Table headings
        table = document.add_table(rows=1, cols=len(algorithms_list))
        hdr_cells = table.rows[0].cells
        for i in range(len(algorithms_list)):
            hdr_cells[i].text = algorithms_list[i]            
           
        for test in keys_exp:    
            #print "new test %s" % test
            for index in range(len(experiments[test])):
                #print "new SUB test %s" % str(index+1)
                row_list = [0.0]*(len(algorithms_list)-1)
                divisor = len(dict_results[key][test][days][str(index+1)])
                for _list in dict_results[key][test][days][str(index+1)]:
                    for model in range(len(_list)):
                        row_list[model] += _list[model]
                for model in range(len(row_list)):
                    row_list[model] = row_list[model]/divisor*100
                ##  header
                row_cells = table.add_row().cells                                           
                row_cells[0].text = test + '.'+ str(index+1)
                col = 1
                ##Table test X_X row
                for cell_list in row_list:
                    cell = ("%.1f%%\n") % (cell_list)
                    row_cells[col].text = cell
                    col += 1

    document.add_page_break()


end = time.time()
print(end - start)

document.save('experiments.docx')
log.close()
print "DONE. Experimets %s" % (num_tests)